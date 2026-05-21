from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(188, 95, 58)
TEXT = RGBColor(33, 40, 51)
MUTED = RGBColor(88, 103, 122)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:{}".format(key)), str(value))


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                left={"val": "nil"},
                bottom={"val": "nil"},
                right={"val": "nil"},
            )


def set_run_font(run, size=None, bold=None, color=None, name="Arial", italic=None):
    run.font.name = name
    rpr = run._element.rPr
    if rpr is not None:
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:hAnsi"), name)
        rfonts.set(qn("w:eastAsia"), name)
        rfonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, label, title):
    label_p = doc.add_paragraph()
    label_p.paragraph_format.space_after = Pt(4)
    label_run = label_p.add_run(label.upper())
    set_run_font(label_run, size=8.5, bold=True, color=ACCENT)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(8)
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=15, bold=True, color=TEXT)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=TEXT)


def add_job(doc, company, role, dates, summary, bullets):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(5.4)
    table.columns[1].width = Inches(1.6)
    remove_table_borders(table)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left_p = left.paragraphs[0]
    left_p.paragraph_format.space_after = Pt(0)
    company_run = left_p.add_run(company)
    set_run_font(company_run, size=12.5, bold=True, color=TEXT)
    role_p = left.add_paragraph()
    role_p.paragraph_format.space_after = Pt(0)
    role_run = role_p.add_run(role)
    set_run_font(role_run, size=10.5, color=MUTED)
    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = right_p.add_run(dates)
    set_run_font(date_run, size=10, bold=True, color=ACCENT)
    doc.add_paragraph()._element.addprevious(table._element)

    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_after = Pt(4)
    summary_p.paragraph_format.line_spacing = 1.18
    summary_run = summary_p.add_run(summary)
    set_run_font(summary_run, size=10.8, color=TEXT)
    for bullet in bullets:
        add_bullet(doc, bullet)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_tag_row(doc, items):
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    remove_table_borders(table)
    for idx, item in enumerate(items):
        cell = table.cell(0, idx)
        shade_cell(cell, "F6EFEA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=9.5, bold=True, color=ACCENT)


def main():
    repo_root = Path(__file__).resolve().parent.parent
    out = repo_root / "assets" / "Yaroslavna-Soldatova-CV.docx"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT

    if "List Bullet" in styles:
        styles["List Bullet"].font.name = "Arial"
        styles["List Bullet"].font.size = Pt(10.5)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run("Yaroslavna Soldatova")
    set_run_font(title_run, size=24, bold=True, color=TEXT)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(6)
    subtitle_run = subtitle_p.add_run("Senior Product Manager / Product Strategist")
    set_run_font(subtitle_run, size=11.5, bold=True, color=ACCENT)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(10)
    meta_text = (
        "Москва, Россия  |  Москва / Remote  |  "
        "yaroslavnaalexandrovna@gmail.com  |  +7 964 879 17 71  |  Telegram: protoptica"
    )
    meta_run = meta_p.add_run(meta_text)
    set_run_font(meta_run, size=10, color=MUTED)

    add_tag_row(
        doc,
        [
            "5+ лет в product management",
            "AI / B2B / сложные системы",
            "Strategy + Roadmap + Execution",
        ],
    )

    profile = (
        "Веду B2B- и AI-продукты в сложной среде: формирую стратегию, расставляю приоритеты, "
        "запускаю новые направления и соединяю продукт, коммерцию и операционные процессы в "
        "одну рабочую систему. Сильнее всего работаю там, где нужно разобраться в сложном "
        "контуре, принять stop/go-решения и довести их до результата."
    )
    profile_p = doc.add_paragraph()
    profile_p.paragraph_format.space_before = Pt(10)
    profile_p.paragraph_format.space_after = Pt(14)
    profile_p.paragraph_format.line_spacing = 1.18
    profile_run = profile_p.add_run(profile)
    set_run_font(profile_run, size=11, color=TEXT)

    add_heading(doc, "Опыт", "Профессиональный опыт")

    add_job(
        doc,
        "Частная стоматология",
        "Product Manager · part-time",
        "04.2026 — н.в.",
        "Веду проект на стыке продуктового аудита, стратегии и legacy-трансформации: от разбора "
        "критического сбоя в коммуникациях до оценки рыночного потенциала платформы.",
        [
            "Перевела проблему с WhatsApp-каналом в измеримую модель потерь: retention, no-show и операционные затраты.",
            "Собрала multi-channel схему коммуникаций: Telegram, MAX, SMS и fallback-сценарии.",
            "Проверила сценарий перехода на SaaS, оценила объем доработок legacy-платформы и подготовила ТЗ с учетом окупаемости.",
            "Снизила риски зависимости от single developer и нестабильных интеграций.",
            "Оценила рентабельность вывода решения на внешний рынок, сроки окупаемости и составила roadmap.",
        ],
    )

    add_job(
        doc,
        "TeamLeaders",
        "Strategic Consultant · part-time",
        "01.2026 — н.в.",
        "Работаю с фаундерами и C-level компаний с оборотом 1 млрд+ над продуктовой стратегией, "
        "приоритетами роста, системой принятия решений и оптимизацией бизнес-процессов.",
        [
            "Собирала из разрозненных инициатив ясный контур решений и приоритетов.",
            "Переводила обсуждения с фаундерами и C-level в гипотезы, аналитические задачи и roadmap.",
            "Помогала командам принимать stop/go-решения на основе данных.",
            "Структурировала стратегию роста через оценку направлений, клиентских сегментов, метрик и ресурсных ограничений.",
        ],
    )

    add_job(
        doc,
        "Третье Мнение",
        "Senior Product Owner",
        "10.2024 — 03.2026",
        "Отвечала за развитие направления, продуктовую стратегию, roadmap, интеграции, UX, "
        "внедрение и техподдержку AI-продуктов для клиник. Работала напрямую с CEO и CPO, принимала "
        "решения по развитию направления и соединяла коммерческую, продуктовую и операционную части "
        "в единый рабочий поток.",
        [
            "Пересобрала один продукт с несколькими модулями в два отдельных продуктовых контура.",
            "Запустила новый AI-модуль с нуля до бета-тестирования за 1 месяц.",
            "Задавала приоритеты и требования для data engineering, чтобы выстроить понятную систему данных, дашбордов и e2e-процессов.",
            "Снизила обращения в поддержку в 5 раз за счет улучшения UX, онбординга и стабильности продукта.",
            "Сократила время поиска нужного видео в 15 раз после перезапуска ключевых UX-сценариев.",
            "Построила модель масштабирования на базе LTV, оттока и конверсии с пилотов.",
            "Остановила одно из нерентабельных направлений видеоаналитики, обосновав решение через рынок, конкуренцию и отсутствие потенциала роста.",
            "Обеспечила рост числа инсталляций в 4 раза.",
        ],
    )

    add_job(
        doc,
        "BIOCAD",
        "Product Owner",
        "07.2017 — 11.2023",
        "Вела R&D-продукты на стыке науки, данных и работы лабораторий: превращала "
        "исследовательские задачи в управляемый продуктовый контур с обоснованными "
        "приоритетами и сроками.",
        [
            "Завершили один из проектов на 3 месяца раньше плана.",
            "Снизила простой лабораторий на 50% благодаря чек-листам и перестройке процессов.",
            "Участвовала в Agile-трансформации департамента R&D, что помогло сократить сроки реализации проектов на 50–60%.",
            "Собрала кросс-функциональную группу по риск-менеджменту и помогла избежать 20% крупных рисков в проектах.",
        ],
    )

    add_job(
        doc,
        "Self-employed",
        "Founder · AI-ассистент для психотерапевтов и психологов",
        "2023 — 2024",
        "Запускала собственный AI-продукт для анализа терапевтических сессий и генерации саммари, "
        "рекомендаций и домашних заданий.",
        [
            "Провела рынок, конкурентный анализ и JTBD-исследование.",
            "Сформировала концепцию MVP, пользовательские сценарии и ключевые метрики.",
            "Проверяла гипотезы через интервью, прототипирование и тестирование с потенциальными пользователями.",
        ],
    )

    add_heading(doc, "Навыки", "Ключевые навыки")
    skills_p = doc.add_paragraph()
    skills_p.paragraph_format.space_after = Pt(10)
    skills_text = (
        "Product strategy, Product management, Prioritization, Roadmap, Discovery, Delivery, JTBD, "
        "Customer Development, UX, Stakeholder management, AI / B2B / legacy-системы, "
        "data-driven decision, риск-менеджмент."
    )
    skills_run = skills_p.add_run(skills_text)
    set_run_font(skills_run, size=10.5, color=TEXT)

    add_heading(doc, "Образование", "Образование")
    edu_items = [
        "Университет современной психологии — психолог-консультант",
        "Межотраслевой институт госаттестации — клинический психолог",
        "Московский государственный университет тонких химических технологий им. М.В. Ломоносова",
        "Саратовский государственный университет им. Н.Г. Чернышевского",
    ]
    for item in edu_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=TEXT)

    add_heading(doc, "Языки", "Языки")
    lang_p = doc.add_paragraph()
    lang_run = lang_p.add_run("Английский — B1 (Intermediate)")
    set_run_font(lang_run, size=10.5, color=TEXT)

    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
